#!/usr/bin/env python3
from docx import Document
import sys

def read_docx(file_path):
    try:
        doc = Document(file_path)
        
        print(f"=== 文件: {file_path} ===")
        print("=" * 80)
        
        # 读取所有段落
        print("\n--- 段落内容 ---")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        # 读取所有表格
        print("\n--- 表格内容 ---")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  行 {row_idx + 1}: {row_data}")
        
        print("\n" + "=" * 80)
        
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python read-docx.py <docx文件路径>")
        sys.exit(1)
    
    read_docx(sys.argv[1])
